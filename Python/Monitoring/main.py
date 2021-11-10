from euipn import EUIPN
import io
import os
import locale
import datetime
import docx
from docx.shared import Cm

LANGUAGE_LOCALE = 'en_US.UTF-8'

def replace_month(paragraphs, month):
    """Add the month name to the document model in the place of the {{Date}} value.

    Args:
        paragraphs (list): A list of paragraph object.
        month (str): The month name.
    """
    for paragraph in paragraphs:
        paragraph.text = paragraph.text.replace(r'{{Date}}', month)

def create_file(file_path, model_file_path, month):
    """Create the file based in the model, name it and create the table and format it.

    Args:
        file_path (str): The path and file name to save the file.
        model_file_path (str): The path to the .docx model file.
        month (str): The month name.

    Returns:
        docx obj: The docx file object.
    """
    print(f'Creating .docx file: {file_path}')
    document = docx.Document(model_file_path)                                                                                                           # the model file should have the table header

    replace_month(document.paragraphs, month)
    table = document.tables[0]                                                                                                                          # get the table object that already exist in the file and fill with the data
    table.autofit = False

    return document


def write_docx(row, document):
    """Fill the .docx model with the data from

    Args:
        data (list): A list of list containing the data to write the .docx file in the order [Figura, Titular, Número, Data Pub., Data Dep., Título].
        docx obj: The docx file object.
    """
    table = document.tables[0]

    new_row = table.add_row()
    new_row.height = Cm(3)

    cells = new_row.cells

    for index, value in zip(range(len(row) + 1), row):                                                                                                  # get each element in the row
        if index == 0:
            cells[index].add_paragraph(value, style='ListNumber')                                                                                       # the first column should be the row index

        elif type(value) == bytes:
            img_cell = cells[index].paragraphs[0].add_run()
            img_cell.add_picture(io.BytesIO(value), height=Cm(2.52))                                                                                    # the image size

        else:
            cells[index].text = value                                                                                                                   # if the value is not a byte like then just writes the value


def get_month_name(date):
    """Get a date and return the month name of the date.

    Args:
        date (ste): The date in the format YYYY-mm-dd to get the month name.

    Returns:
        str: The month name
    """
    locale.setlocale(locale.LC_ALL, LANGUAGE_LOCALE)                                                                                                    # the language to the month name to be displayed

    month_name = datetime.datetime.strptime(date[5:7], "%m").strftime("%B")
    return month_name.capitalize()

def main():
    path = os.path.dirname(os.path.realpath(__file__))                                                                                                  # get this file path, this should be the root path to the others files that will be used
    start_date, end_date = EUIPN.get_dates()                                                                                                            # get the first day from the last month and the last day from the last month
    month = get_month_name(start_date)

    auth = {'username': 'user',
            'password':'password'}

    euipn = EUIPN(**auth)

    replaces = {'"start_date"': f'"{start_date}"', '"end_date"': f'"{end_date}"'}
    query_parameters = euipn.get_query_parameters(f'{path}\query_parameters.json', replaces)                                                            # get the query paramethers from the json file and replace the first data and the end date with the value

    default_file_path = '{path}\\files\Monitoring ({region}) - {month}.{year}.docx'                                                                     # the file name pattern
    model_file_path =f'{path}\Model.docx'                                                                                                               # the file model path, this file has the font style and table style


    columns = ['', 'preferredImageUrl', 'applicantName', 'designNumber', 'publicationDate', 'registrationDate', 'indicationOfProduct']

    for query_parameter in query_parameters:
        region = query_parameter.pop('fileName')
        file_path = default_file_path.format(path=path, region=region, month=month[:3], year=start_date[:4])
        document = create_file(file_path, model_file_path, month)

        for row in euipn.get_design_data(query_parameter, columns):
            write_docx(row, document)

            document.save(file_path)

    print('Process finished!')

main()
